#!/usr/bin/env python3
"""
tex2docx.py  --  Overleaf LaTeX  ->  clean, editable Word (.docx)

One command turns a paper written in Overleaf (LaTeX + figures + a
Zotero-synced references.bib) into a Word document that a non-LaTeX reader
can open, read, and comment on.  It keeps the fonts (Times New Roman, 12 pt,
single spacing), real Word equations, the figures, and a proper ACS-style
reference list generated straight from the .bib.

    python tex2docx.py main.tex
    python tex2docx.py main.tex -o Schauer_QE.docx
    python tex2docx.py main.tex --keep-drafts --pdf

Requirements on your machine:
    * pandoc            (brew install pandoc)          -- does the heavy lifting
    * internet, once    -- to fetch the ACS citation style the first time
    * (optional) LibreOffice for --pdf preview; a PDF-figure converter
      (macOS 'sips' is built in) if your figures are .pdf/.eps

Everything else (the reference.docx template, the ACS .csl) lives next to
this script and is reused every run.
"""

import argparse
import os
import re
import shutil
import subprocess
import sys
import urllib.request

from .constants import (PAGEBREAK_MARK, APPENDIX_MARK, REFERENCES_MARK,
                        XREF_A, XREF_S, XREF_E,
                        REFERENCE_DOCX, CSL_PATH, CSL_URL)
from .counters import (scan_preamble, build_label_map, float_numbers,
                       expand_counter_macros, strip_comments)
from . import colors as _colors


# ----------------------------------------------------------------------------
# small helpers
# ----------------------------------------------------------------------------

SUB = str.maketrans("0123456789+-=()", "\u2080\u2081\u2082\u2083\u2084"
                    "\u2085\u2086\u2087\u2088\u2089\u208a\u208b\u208c\u208d\u208e")
SUP = str.maketrans("0123456789+-=()n", "\u2070\u00b9\u00b2\u00b3\u2074"
                    "\u2075\u2076\u2077\u2078\u2079\u207a\u207b\u207c\u207d\u207e\u207f")


def log(msg):
    print(f"  {msg}")


def find_bib(tex_src, tex_dir):
    """Locate the .bib: honour \\addbibresource / \\bibliography, else look for
    a single .bib next to the .tex."""
    m = re.search(r"\\(?:addbibresource|bibliography)\{([^}]+)\}", tex_src)
    if m:
        name = m.group(1).strip()
        if not name.lower().endswith(".bib"):
            name += ".bib"
        cand = os.path.join(tex_dir, name)
        if os.path.exists(cand):
            return cand
    bibs = [f for f in os.listdir(tex_dir) if f.lower().endswith(".bib")]
    if len(bibs) == 1:
        return os.path.join(tex_dir, bibs[0])
    if bibs:
        log(f"note: several .bib files found, using {bibs[0]}")
        return os.path.join(tex_dir, bibs[0])
    return None


def parse_width(opts):
    """Interpret an \\includegraphics option string.

    Returns ("frac", 0.85) for width=0.85\\textwidth, ("abs", inches) for an
    absolute unit like width=2in, or None when no width is given.
    """
    opts = opts or ""
    m = re.search(r"width\s*=\s*([0-9]*\.?[0-9]+)?\s*\\(?:text|line|column)width",
                  opts)
    if m:
        return ("frac", float(m.group(1)) if m.group(1) else 1.0)
    m = re.search(r"width\s*=\s*([0-9]*\.?[0-9]+)\s*(in|cm|mm|pt|bp)\b", opts)
    if m:
        value = float(m.group(1))
        per_inch = {"in": 1.0, "cm": 2.54, "mm": 25.4, "pt": 72.27, "bp": 72.0}
        return ("abs", value / per_inch[m.group(2)])
    return None


def scan_floats(src):
    """Walk the document once and return, in document order:
       * img_widths  : one width-fraction (or None) per \\includegraphics
       * fig_labels  : one \\label (or None) per figure environment
       * tab_labels  : one \\label (or None) per table environment
    Used to resize/centre images and to number + bookmark captions."""
    img_widths, fig_labels, tab_labels = [], [], []
    stack = []
    tok = re.compile(
        r"\\begin\{(figure|table)\}\*?|"
        r"\\end\{(figure|table)\}\*?|"
        r"\\includegraphics\s*(\[[^\]]*\])?\s*\{[^}]*\}|"
        r"\\label\{([^}]+)\}")
    for m in tok.finditer(src):
        if m.group(1):                       # \begin{figure|table}
            stack.append(m.group(1))
            (fig_labels if m.group(1) == "figure" else tab_labels).append(None)
        elif m.group(2):                     # \end{figure|table}
            if stack:
                stack.pop()
        elif m.group(0).startswith("\\includegraphics"):
            img_widths.append(parse_width(m.group(3)))
        elif m.group(4):                     # \label{...}
            if stack and stack[-1] == "figure" and fig_labels and fig_labels[-1] is None:
                fig_labels[-1] = m.group(4)
            elif stack and stack[-1] == "table" and tab_labels and tab_labels[-1] is None:
                tab_labels[-1] = m.group(4)
    return img_widths, fig_labels, tab_labels


def collect_referenced(src):
    """Every label targeted by a \\cref/\\Cref/\\autoref/\\ref, so we know which
    bookmarks to keep as link targets."""
    out = set()
    for m in re.finditer(r"\\(?:Cref|cref|autoref|ref)\{([^}]*)\}", src):
        for p in m.group(1).split(","):
            if p.strip():
                out.add(p.strip())
    return out


def parse_ch(body):
    """Render a chemformula/mhchem argument (e.g. 'MgCl2', 'Mg^2+',
    '(NH2)2CO', 'Ca^2+') to Unicode with sub/superscripts."""
    out = []
    i = 0
    while i < len(body):
        c = body[i]
        if c == "^":                      # charge / superscript: ^2+  ^{2+}
            i += 1
            if i < len(body) and body[i] == "{":
                j = body.index("}", i)
                grp = body[i + 1:j]
                i = j + 1
            else:
                grp = ""
                while i < len(body) and body[i] in "0123456789+-":
                    grp += body[i]
                    i += 1
            out.append(grp.translate(SUP))
        elif c == "_":                    # explicit subscript _2 or _{2}
            i += 1
            if i < len(body) and body[i] == "{":
                j = body.index("}", i)
                grp = body[i + 1:j]
                i = j + 1
            else:
                grp = body[i] if i < len(body) else ""
                i += 1
            out.append(grp.translate(SUB))
        elif c.isdigit():                 # digit right after a letter -> subscript
            if out and (out[-1][-1:].isalpha() or out[-1][-1:] == ")"):
                out.append(c.translate(SUB))
            else:
                out.append(c)
            i += 1
        elif c in "+-" and (i + 1 >= len(body) or body[i + 1] in " \t)}>,;"):
            out.append(c.translate(SUP))   # bare trailing charge, e.g. K+ -> K⁺
            i += 1
        else:
            out.append(c)
            i += 1
    return "".join(out)


def replace_braced(src, macro, handler):
    """Replace every \\macro{...} (brace-matched) using handler(inner)->str."""
    token = "\\" + macro + "{"
    out = []
    i = 0
    while True:
        k = src.find(token, i)
        if k == -1:
            out.append(src[i:])
            break
        out.append(src[i:k])
        j = k + len(token)
        depth = 1
        while j < len(src) and depth:
            if src[j] == "{":
                depth += 1
            elif src[j] == "}":
                depth -= 1
            j += 1
        inner = src[k + len(token):j - 1]
        out.append(handler(inner))
        i = j
    return "".join(out)


def convert_pdf_figure(path):
    """Best-effort convert a .pdf/.eps figure to .png so Word can embed it.
    Returns the new path, or None if no converter is available."""
    png = os.path.splitext(path)[0] + "_conv.png"
    if os.path.exists(png):
        return png
    for cmd in (["sips", "-s", "format", "png", path, "--out", png],          # macOS
                ["magick", "-density", "200", path, png],                      # ImageMagick 7
                ["convert", "-density", "200", path, png],                     # ImageMagick 6
                ["pdftoppm", "-png", "-r", "200", "-singlefile", path,
                 os.path.splitext(png)[0]]):                                   # poppler
        if shutil.which(cmd[0]):
            try:
                subprocess.run(cmd, check=True, capture_output=True)
                if os.path.exists(png):
                    return png
            except subprocess.CalledProcessError:
                continue
    return None


# ----------------------------------------------------------------------------
# preprocessing: make the .tex friendlier for pandoc
# ----------------------------------------------------------------------------

def strip_newcommand(src, name):
    """Remove a \\newcommand/\\renewcommand/\\providecommand definition of \\name
    (with optional [n] arg spec and a brace-matched body).  We do this for the
    few macros we rewrite ourselves, so our text substitution never corrupts
    the macro's own definition line."""
    for kw in ("newcommand", "renewcommand", "providecommand"):
        anchor = "\\" + kw + "{\\" + name + "}"
        k = src.find(anchor)
        if k == -1:
            continue
        j = k + len(anchor)
        while j < len(src) and src[j] in " \t":
            j += 1
        if j < len(src) and src[j] == "[":            # optional [n]
            j = src.index("]", j) + 1
        while j < len(src) and src[j] in " \t":
            j += 1
        if j < len(src) and src[j] == "{":            # brace-matched body
            depth = 1
            j += 1
            while j < len(src) and depth:
                if src[j] == "{":
                    depth += 1
                elif src[j] == "}":
                    depth -= 1
                j += 1
        src = src[:k] + src[j:]
    return src


# ---------------------------------------------------------------------------
# figure resolution
# ---------------------------------------------------------------------------

def make_resolver(src, tex_dir):
    """Return resolve(name) -> path, mimicking LaTeX's own figure search.

    LaTeX looks in every \\graphicspath directory and appends the extension
    itself, so `\\includegraphics{Fig2}` legitimately means `Figures/Fig2.pdf`.
    Resolving only literal ".pdf" names in the main directory -- as OWL used to
    -- silently missed every figure in a document that uses either feature.
    """
    search_dirs = ["."]
    for gp in re.findall(r"\\graphicspath\s*\{((?:\s*\{[^}]*\}\s*)+)\}", src):
        for d in re.findall(r"\{([^}]*)\}", gp):
            d = d.strip()
            if d and d not in search_dirs:
                search_dirs.append(d)
    for extra in ("Figures", "figures", "Figure", "figure", "img", "images"):
        if extra not in search_dirs:
            search_dirs.append(extra)

    TRY_EXT = ("", ".pdf", ".png", ".jpg", ".jpeg", ".eps",
               ".PDF", ".PNG", ".JPG")

    def resolve(name):
        name = name.strip()
        for d in search_dirs:
            for ext in TRY_EXT:
                cand = os.path.normpath(os.path.join(tex_dir, d, name + ext))
                if os.path.isfile(cand):
                    return cand
        return None

    resolve.search_dirs = search_dirs
    return resolve


def preprocess(src, tex_dir, labels, keep_drafts):
    """Returns (preprocessed_source, img_widths)."""
    # 0) comments first: a commented-out \\includegraphics or \\appsection
    #    must not be counted as real content
    src = strip_comments(src)

    # 0a) expand user macros that step a counter (e.g. \appsection) -- pandoc
    #     does not run TeX, so \theappx would otherwise vanish and the heading
    #     would read "Appendix : Title" with no letter
    src = expand_counter_macros(src)

    # 0) drop the definitions of the macros we handle ourselves, so our text
    #    substitutions below can't damage their own \newcommand line
    for name in ("outlinetag", "draftnote"):
        src = strip_newcommand(src, name)

    # 0b) unwrap spacing environments (setspace): pandoc doesn't know them and
    #     would DROP their entire contents -- which is where the title page lives
    src = re.sub(r"\\begin\{(?:singlespace|singlespacing|doublespace|"
                 r"doublespacing|onehalfspace|onehalfspacing)\}", "", src)
    src = re.sub(r"\\end\{(?:singlespace|singlespacing|doublespace|"
                 r"doublespacing|onehalfspace|onehalfspacing)\}", "", src)
    src = re.sub(r"\\begin\{spacing\}\{[^}]*\}", "", src)
    src = re.sub(r"\\end\{spacing\}", "", src)

    # 0c) the author's paragraph style is "\\" then "\noindent"; turn that into
    #     a real paragraph break so justified text doesn't stretch the last line
    src = re.sub(r"\\\\\s*\n?\s*\\noindent\b", "\n\n", src)
    src = re.sub(r"\\noindent\b", "", src)

    # 1) chemistry: \ch{...} and \ce{...} -> Unicode
    src = replace_braced(src, "ch", parse_ch)
    src = replace_braced(src, "ce", parse_ch)

    # 2) draft notes / outline banners
    if keep_drafts:
        src = replace_braced(src, "draftnote",
                             lambda s: r"\textbf{[Draft note]} " + s)
    else:
        src = replace_braced(src, "draftnote", lambda s: "")
    src = src.replace(r"\outlinetag", "")

    # 3) cross-references: \cref / \Cref / \autoref / \ref  ->  clickable
    #    "Section 2" / "Figure 1" links (token now, real hyperlink in postprocess)
    def xref(inner):
        parts = [p.strip() for p in inner.split(",") if p.strip()]
        names = []
        for p in parts:
            kind, num = labels.get(p, (None, None))
            if kind and num:
                names.append(XREF_A + p + XREF_S + f"{kind} {num}" + XREF_E)
            elif kind:
                names.append(kind)
            else:
                names.append(p)
        if len(names) <= 1:
            return names[0] if names else ""
        return ", ".join(names[:-1]) + " and " + names[-1]
    for mac in ("Cref", "cref", "autoref", "ref"):
        src = replace_braced(src, mac, xref)

    # 3b) section boundaries first (they carry a page-number-format change), then
    #     ordinary page breaks. Distinct sentinels so the post-processor can make
    #     the back matter its own section(s) with Roman numerals + new header.
    src = re.sub(r"\\clearpage\s*\\pagenumbering\{Roman\}",
                 "\n\n" + APPENDIX_MARK + "\n\n", src)
    src = re.sub(r"\\clearpage\s*\\fancyhead\[L\]\{[^}]*References[^}]*\}",
                 "\n\n" + REFERENCES_MARK + "\n\n", src)
    src = re.sub(r"\\(?:newpage|clearpage|pagebreak)\b",
                 "\n\n" + PAGEBREAK_MARK + "\n\n", src)
    # strip page-style commands so they don't leak into the text
    src = re.sub(r"\\pagenumbering\{[^}]*\}", "", src)
    src = re.sub(r"\\fancyhead\[[^\]]*\]\{[^}]*\}", "", src)
    src = re.sub(r"\\(?:thispagestyle|pagestyle)\{[^}]*\}", "", src)

    # 3b) \textcolor -> sentinels the postprocessor turns into real colour
    src = _colors.mark(src)

    # 4) resolve figure paths and convert .pdf / .eps so Word can embed them
    resolve_graphic = make_resolver(src, tex_dir)

    def fix_graphic(name, opts):
        """Rewrite one \\includegraphics and, when it resolves, record its
        requested width.

        The width list has to be built HERE rather than by scanning the source
        separately: a commented-out example in the preamble, or a figure file
        that is missing, changes how many images pandoc actually emits, and a
        positional list built from the raw source then applies every width to
        the wrong picture.  Collecting them alongside the rewrite keeps the
        list exactly in step with the images in the .docx.
        """
        name = name.strip()
        path = resolve_graphic(name)
        if not path:
            log(f"note: figure '{name}' not found")
            return name
        img_widths.append(parse_width(opts))
        if path.lower().endswith((".pdf", ".eps")):
            conv = convert_pdf_figure(path)
            if conv:
                return os.path.relpath(conv, tex_dir)
            log(f"note: no PDF->PNG converter for '{name}'; "
                f"install poppler-utils or ImageMagick, or supply a .png")
            return name
        return os.path.relpath(path, tex_dir)

    # NB: must handle the optional argument too -- \includegraphics[width=...]{f}
    # is the normal form, and a plain brace-match would skip every one of them.
    img_widths = []
    src = re.sub(r"\\includegraphics(\s*\[[^\]]*\])?\s*\{([^}]*)\}",
                 lambda m: "\\includegraphics" + (m.group(1) or "") +
                           "{" + fix_graphic(m.group(2), m.group(1)) + "}", src)

    return src, img_widths


# ----------------------------------------------------------------------------
# assets: ACS csl (download once) + reference.docx (built by make_reference.py)
# ----------------------------------------------------------------------------

def postprocess_docx(path, header_title, floats=((), (), ()), referenced=frozenset(),
                     numbering=((), ())):
    """After pandoc, rebuild the page furniture so it matches the LaTeX:
      * centre + enlarge the cover block (title page);
      * resize images to the requested fraction of the text width and centre them;
      * number captions ("Figure N:" / "Table N:") and bookmark them;
      * turn \\cref/\\ref cross-references into clickable internal hyperlinks;
      * turn \\newpage sentinels into real Word page breaks;
      * split the back matter into its own section(s) with a running header
        and Roman page numbers, the body with Arabic numbers starting at 1,
        and the cover with no header/number.
    Needs python-docx; without it, falls back to plain page breaks."""
    try:
        import copy
        import docx
        from docx.shared import Pt, Emu
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError:
        _pagebreaks_rawxml(path)
        log("note: install python-docx for the title page, running header, "
            "and page numbering (pip install python-docx)")
        return

    d = docx.Document(path)

    def clear_runs(p):
        for r in list(p.runs):
            r._element.getparent().remove(r._element)

    def page_field(p):
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        r = OxmlElement("w:r"); t = OxmlElement("w:t"); t.text = "1"
        r.append(t); fld.append(r)
        p._p.append(fld)

    def set_pgnum(section, fmt=None, start=None):
        sectPr = section._sectPr
        for e in sectPr.findall(qn("w:pgNumType")):
            sectPr.remove(e)
        el = OxmlElement("w:pgNumType")
        if fmt:
            el.set(qn("w:fmt"), fmt)
        if start is not None:
            el.set(qn("w:start"), str(start))
        sectPr.append(el)

    def build_header(section, left):
        section.header.is_linked_to_previous = False
        h = section.header.paragraphs[0]
        clear_runs(h)
        h.text = ""
        from docx.shared import Inches
        pw = section.page_width or Inches(8.5)
        lm = section.left_margin or Inches(1)
        rm = section.right_margin or Inches(1)
        h.paragraph_format.tab_stops.add_tab_stop(pw - lm - rm,
                                                  WD_TAB_ALIGNMENT.RIGHT)
        run = h.add_run(left)
        run.italic = True
        run.font.size = Pt(10)
        tab = h.add_run("\t"); tab.font.size = Pt(10)
        page_field(h)
        # thin rule under the header, like fancyhdr's headrule
        pPr = h._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        for k, v in (("w:val", "single"), ("w:sz", "4"),
                     ("w:space", "1"), ("w:color", "auto")):
            bottom.set(qn(k), v)
        pbdr.append(bottom); pPr.append(pbdr)

    # ── figures / tables / cross-references ──────────────────────────────
    img_widths, fig_labels, tab_labels = floats
    fig_nums, tab_nums = numbering
    _bid = [9000]

    # 1) figures: pandoc wraps each figure in a 2-column "FigureTable" (which
    #    jams the image into one cell -> off-centre and size-constrained). We
    #    UNWRAP every figure image into a clean centred body paragraph sized to
    #    the requested fraction of the text width. Inline images that aren't in a
    #    figure table (e.g. the title crest) are just resized in place + centred.
    import io as _io

    def _ancestor(el, tag):
        while el is not None and el.tag != qn(tag):
            el = el.getparent()
        return el

    def _center(pel):
        pPr = pel.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr"); pel.insert(0, pPr)
        for j in pPr.findall(qn("w:jc")):
            pPr.remove(j)
        jc = OxmlElement("w:jc"); jc.set(qn("w:val"), "center"); pPr.append(jc)

    sec0 = d.sections[0]
    text_w = ((sec0.page_width or Emu(7772400))
              - (sec0.left_margin or Emu(914400))
              - (sec0.right_margin or Emu(914400)))
    img_i = 0
    for draw in list(d.element.iter(qn("w:drawing"))):
        spec = img_widths[img_i] if img_i < len(img_widths) else None
        frac = spec[1] if spec and spec[0] == "frac" else None
        abs_in = spec[1] if spec and spec[0] == "abs" else None
        img_i += 1
        ext = draw.find(".//" + qn("wp:extent"))
        ncx = int(ext.get("cx")) if ext is not None else 0
        if frac:
            target = int(text_w * frac)
        elif abs_in:
            target = int(abs_in * 914400)
        else:
            target = ncx or None
        tbl = _ancestor(draw, "w:tbl")
        if tbl is not None:                           # figure wrapped in a table
            blip = draw.find(".//" + qn("a:blip"))
            rid = blip.get(qn("r:embed")) if blip is not None else None
            try:
                blob = d.part.related_parts[rid].blob if rid else None
            except Exception:
                blob = None
            if not blob or not target:
                continue
            try:
                d.add_picture(_io.BytesIO(blob), width=Emu(target))
            except Exception:
                continue
            new_p = d.paragraphs[-1]._p                # the paragraph add_picture made
            _center(new_p)
            tbl.addprevious(new_p)                     # move it in front of the table
            tbl.getparent().remove(tbl)                # drop the wrapper table
        else:                                          # inline image (not a figure)
            pel = _ancestor(draw, "w:p")
            if pel is not None:
                _center(pel)
            if target and target != ncx and ext is not None and ncx:
                cy = int(ext.get("cy"))
                for e in (ext, draw.find(".//" + qn("a:ext"))):
                    if e is not None:
                        e.set("cx", str(target)); e.set("cy", str(int(cy * target / ncx)))

    # 2) number captions and bookmark them for cross-references
    def _norm(s):
        return (s or "").replace(" ", "").lower()

    def _prepend_bold(p, text):
        r = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr"); rpr.append(OxmlElement("w:b")); r.append(rpr)
        t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
        r.append(t)
        pPr = p._p.find(qn("w:pPr"))
        (pPr.addnext(r) if pPr is not None else p._p.insert(0, r))

    def _bookmark(p, name):
        _bid[0] += 1
        s = OxmlElement("w:bookmarkStart")
        s.set(qn("w:id"), str(_bid[0])); s.set(qn("w:name"), name)
        e = OxmlElement("w:bookmarkEnd"); e.set(qn("w:id"), str(_bid[0]))
        pPr = p._p.find(qn("w:pPr"))
        (pPr.addnext(s) if pPr is not None else p._p.insert(0, s))
        p._p.append(e)

    # Caption numbers come from the same counter machine the cross-references
    # use, so a supplementary section that restarts the counter as
    # \thefigure = S\arabic{figure} yields "Figure S1", not "Figure 6", and
    # the caption agrees with every \cref pointing at it.
    fig_i = tab_i = 0
    for p in d.paragraphs:
        st = _norm(p.style.name if p.style else "")
        if st == "imagecaption":
            num = fig_nums[fig_i] if fig_i < len(fig_nums) else str(fig_i + 1)
            _prepend_bold(p, f"Figure {num}: ")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            lbl = fig_labels[fig_i] if fig_i < len(fig_labels) else None
            if lbl:
                _bookmark(p, lbl)
            fig_i += 1
        elif st == "tablecaption":
            num = tab_nums[tab_i] if tab_i < len(tab_nums) else str(tab_i + 1)
            _prepend_bold(p, f"Table {num}: ")
            lbl = tab_labels[tab_i] if tab_i < len(tab_labels) else None
            if lbl:
                _bookmark(p, lbl)
            tab_i += 1

    # 3) cross-reference tokens -> real clickable internal hyperlinks
    tokre = re.compile(re.escape(XREF_A) + r"(.*?)" + re.escape(XREF_S)
                       + r"(.*?)" + re.escape(XREF_E))

    def _run(text, rpr):
        r = OxmlElement("w:r")
        if rpr is not None:
            r.append(copy.deepcopy(rpr))
        t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
        r.append(t)
        return r

    def _link(anchor, disp, rpr):
        hl = OxmlElement("w:hyperlink"); hl.set(qn("w:anchor"), anchor)
        hl.append(_run(disp, rpr))
        return hl

    for r in list(d.element.iter(qn("w:r"))):
        ts = r.findall(qn("w:t"))
        if not ts:
            continue
        txt = "".join(t.text or "" for t in ts)
        if XREF_A not in txt:
            continue
        rpr = r.find(qn("w:rPr"))
        parent = r.getparent(); idx = list(parent).index(r)
        nodes, pos = [], 0
        for m in tokre.finditer(txt):
            if m.start() > pos:
                nodes.append(_run(txt[pos:m.start()], rpr))
            nodes.append(_link(m.group(1), m.group(2), rpr))
            pos = m.end()
        if pos < len(txt):
            nodes.append(_run(txt[pos:], rpr))
        parent.remove(r)
        for k, node in enumerate(nodes):
            parent.insert(idx + k, node)

    paras = d.paragraphs
    base_sectPr = d.sections[-1]._sectPr

    # --- cover block: everything before the first (any) sentinel -> centred ---
    def is_mark(p):
        return any(m in p.text for m in
                   (PAGEBREAK_MARK, APPENDIX_MARK, REFERENCES_MARK))
    first = next((i for i, p in enumerate(paras) if is_mark(p)), None)
    if first is not None:
        title_done = False
        for p in paras[:first]:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if not title_done and p.text.strip():
                for r in p.runs:
                    r.font.size = Pt(19); r.font.bold = True
                title_done = True

    # --- insert section breaks at the back-matter sentinels ---
    def make_section_break(p):
        pPr = p._p.get_or_add_pPr()
        new = copy.deepcopy(base_sectPr)
        for tag in ("w:headerReference", "w:footerReference",
                    "w:pgNumType", "w:titlePg"):
            for e in new.findall(qn(tag)):
                new.remove(e)
        pPr.append(new)
        clear_runs(p)

    for p in paras:
        if APPENDIX_MARK in p.text or REFERENCES_MARK in p.text:
            make_section_break(p)

    # --- ordinary page breaks ---
    for p in paras:
        if PAGEBREAK_MARK in p.text:
            clear_runs(p)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.add_run().add_break(WD_BREAK.PAGE)

    # --- configure each section: header + numbering ---
    secs = d.sections
    # remove the heading bookmarks (the grey [ ] markers). Pandoc bookmarks every
    # heading; keep the "ref-..." citation targets AND any label a cross-reference
    # actually points at (so the clickable links resolve).
    remove_ids = set()
    for bm in d.element.iter(qn("w:bookmarkStart")):
        name = bm.get(qn("w:name")) or ""
        if name.startswith("ref-") or name in referenced:
            continue
        remove_ids.add(bm.get(qn("w:id")))
    for tag in ("w:bookmarkStart", "w:bookmarkEnd"):
        for bm in list(d.element.iter(qn(tag))):
            if bm.get(qn("w:id")) in remove_ids:
                bm.getparent().remove(bm)

    # sec[0] = cover + body ; then appendix ; then references (if present)
    body = secs[0]
    body.different_first_page_header_footer = True     # cover page: no header
    build_header(body, header_title)
    set_pgnum(body, fmt="decimal", start=0)            # cover=0 (hidden), body=1..

    labels = ["Appendix", "References"]
    for sec, label in zip(secs[1:], labels):
        sec.different_first_page_header_footer = False
        build_header(sec, label)
        # first back-matter section restarts Roman at I; later ones continue
        if label == labels[0]:
            set_pgnum(sec, fmt="upperRoman", start=1)
        else:
            set_pgnum(sec, fmt="upperRoman")

    _colors.apply(d)

    d.save(path)


def _pagebreaks_rawxml(path):
    """Fallback: swap sentinel paragraphs for page breaks by editing the XML."""
    import zipfile
    import tempfile
    tmp = path + ".tmp"
    with zipfile.ZipFile(path) as zin:
        names = zin.namelist()
        data = {n: zin.read(n) for n in names}
    xml = data["word/document.xml"].decode("utf-8")
    xml = re.sub(
        r"<w:p\b[^>]*>(?:(?!</w:p>).)*?" + PAGEBREAK_MARK + r".*?</w:p>",
        '<w:p><w:r><w:br w:type="page"/></w:r></w:p>', xml, flags=re.S)
    data["word/document.xml"] = xml.encode("utf-8")
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for n in names:
            zout.writestr(n, data[n])
    os.replace(tmp, path)


def ensure_csl():
    if os.path.exists(CSL_PATH):
        return CSL_PATH
    log("fetching the ACS citation style (one-time)...")
    try:
        req = urllib.request.Request(CSL_URL, headers={"User-Agent": "tex2docx"})
        with urllib.request.urlopen(req, timeout=30) as r:
            data = r.read()
        with open(CSL_PATH, "wb") as f:
            f.write(data)
        return CSL_PATH
    except Exception as e:
        log(f"could not download the ACS style ({e}).")
        log("  -> the doc will still build with pandoc's default numeric style.")
        log("  -> to match ACS exactly, drop american-chemical-society.csl next "
            "to this script")
        log("     (Zotero already has it: Preferences > Cite > Styles).")
        return None


# ----------------------------------------------------------------------------
# main
# ----------------------------------------------------------------------------

def _bib_field(entry, field):
    """Return the brace/quote-balanced value of a bib field, or ''."""
    m = re.search(r"(?<![A-Za-z])" + field + r"\s*=\s*", entry, re.I)
    if not m:
        return ""
    i = m.end()
    if i >= len(entry):
        return ""
    if entry[i] == "{":
        depth, j = 0, i
        while j < len(entry):
            if entry[j] == "{":
                depth += 1
            elif entry[j] == "}":
                depth -= 1
                if depth == 0:
                    return entry[i + 1:j]
            j += 1
        return ""
    if entry[i] == '"':
        j = entry.find('"', i + 1)
        return entry[i + 1:j] if j > 0 else ""
    m2 = re.match(r"([^,\n]+)", entry[i:])
    return m2.group(1).strip() if m2 else ""


def bib_author_year_map(path):
    """key -> (first-author surname, year, has_multiple_authors)."""
    with open(path, encoding="utf-8", errors="ignore") as f:
        text = f.read()
    out = {}
    for em in re.finditer(r"@\w+\s*\{", text):
        start = em.end() - 1
        depth, j = 0, start
        while j < len(text):
            if text[j] == "{":
                depth += 1
            elif text[j] == "}":
                depth -= 1
                if depth == 0:
                    break
            j += 1
        body = text[start + 1:j]
        key = body.split(",", 1)[0].strip()
        if not key:
            continue
        author = _bib_field(body, "author") or _bib_field(body, "editor")
        date = _bib_field(body, "date") or _bib_field(body, "year")
        ym = re.search(r"\d{4}", date)
        year = ym.group(0) if ym else ""
        first = re.split(r"\s+and\s+", author.strip())[0].replace("{", "").replace("}", "")
        if "," in first:
            surname = first.split(",")[0].strip()
        else:
            surname = first.split()[-1] if first.split() else first
        out[key] = (surname.strip(), year, " and " in author)
    return out


def to_zotero_markers(src, cmap):
    """Replace \\cite{...} with Zotero RTF/ODF-Scan markers {Author, Year} and
    swap \\printbibliography for a {Bibliography} placeholder."""
    def one(key):
        if key in cmap:
            sur, yr, many = cmap[key]
            label = sur + (" et al." if many else "")
            return f"{label}, {yr}" if yr else label
        return key

    def repl(m):
        keys = [k.strip() for k in m.group(1).split(",") if k.strip()]
        return "\\{" + "; ".join(one(k) for k in keys) + "\\}"

    src = re.sub(r"\\(?:cite|parencite|autocite|citep|citet|textcite)\s*"
                 r"(?:\[[^\]]*\])?\s*\{([^}]*)\}", repl, src)
    src = re.sub(r"\\printbibliography(?:\[[^\]]*\])?",
                 lambda m: "\\section*{References}\n\n\\{Bibliography\\}", src)
    return src


def main():
    ap = argparse.ArgumentParser(description="Overleaf LaTeX -> editable Word")
    ap.add_argument("tex", help="the .tex file (e.g. main.tex)")
    ap.add_argument("-o", "--output", help="output .docx (default: same name)")
    ap.add_argument("--keep-drafts", action="store_true",
                    help="keep \\draftnote{...} notes instead of removing them")
    ap.add_argument("--pdf", action="store_true",
                    help="also render a .pdf preview via LibreOffice")
    ap.add_argument("--zotero", action="store_true",
                    help="emit citations as Zotero RTF/ODF-Scan markers so they "
                         "can be turned into live, editable Zotero fields")
    args = ap.parse_args()

    if not shutil.which("pandoc"):
        sys.exit("ERROR: pandoc is not installed.  ->  brew install pandoc")

    tex_path = os.path.abspath(args.tex)
    if not os.path.exists(tex_path):
        sys.exit(f"ERROR: no such file: {tex_path}")
    tex_dir = os.path.dirname(tex_path)
    suffix = "_zotero.docx" if args.zotero else ".docx"
    out_path = os.path.abspath(args.output) if args.output else \
        os.path.splitext(tex_path)[0] + suffix

    print(f"tex2docx: {os.path.basename(tex_path)} -> {os.path.basename(out_path)}")

    with open(tex_path, encoding="utf-8") as f:
        src = f.read()

    bib = find_bib(src, tex_dir)
    if bib:
        log(f"bibliography: {os.path.basename(bib)}")
    else:
        log("no .bib found -- citations will be left as-is")

    # running-header text: the \shorttitle macro if present, else the title
    m = re.search(r"\\newcommand\{\\shorttitle\}\{(.+?)\}", src)
    if not m:
        m = re.search(r"\\newcommand\{\\reporttitle\}\{(.+?)\}", src)
    header_title = m.group(1) if m else os.path.splitext(os.path.basename(tex_path))[0]
    header_title = (header_title.replace("---", "\u2014").replace("--", "\u2013")
                    .replace("\\&", "&").strip())

    model = scan_preamble(src)
    labels = build_label_map(src, model)
    fig_nums, tab_nums = float_numbers(src, scan_preamble(src))
    floats = scan_floats(src)              # image widths + figure/table labels
    referenced = collect_referenced(src)   # labels targeted by cross-references
    src = preprocess(src, tex_dir, labels, args.keep_drafts)

    if args.zotero:
        if not bib:
            sys.exit("ERROR: --zotero needs a .bib (none found).")
        src = to_zotero_markers(src, bib_author_year_map(bib))
        log("citations emitted as Zotero RTF/ODF-Scan markers")

    pre_path = os.path.join(tex_dir, ".tex2docx.pre.tex")
    with open(pre_path, "w", encoding="utf-8") as f:
        f.write(src)

    cmd = ["pandoc", pre_path, "-o", out_path,
           "--from", "latex+raw_tex",
           "--reference-doc", REFERENCE_DOCX if os.path.exists(REFERENCE_DOCX) else None,
           "--resource-path", tex_dir]
    cmd = [c for c in cmd if c is not None]

    if bib and not args.zotero:
        csl = ensure_csl()
        cmd += ["--citeproc", "--bibliography", bib,
                "--metadata", "reference-section-title=References",
                "--metadata", "link-citations=true"]
        if csl:
            cmd += ["--csl", csl]

    log("running pandoc...")
    proc = subprocess.run(cmd, capture_output=True, text=True)
    if proc.returncode != 0:
        print(proc.stderr)
        os.remove(pre_path)
        sys.exit("ERROR: pandoc failed (see message above).")
    if proc.stderr.strip():
        for line in proc.stderr.strip().splitlines():
            log(f"pandoc: {line}")
    os.remove(pre_path)
    postprocess_docx(out_path, header_title, floats, referenced,
                     numbering=(fig_nums, tab_nums))
    print(f"OK  ->  {out_path}")
    if args.zotero:
        print("  -> citations are Zotero scan markers {Author, Year}. To make "
              "them live:\n"
              "     1. import your .bib into Zotero (File > Import)\n"
              "     2. open this .docx in LibreOffice and Save As .odt\n"
              "     3. Zotero > Tools > RTF/ODF Scan > ODF (to ODF) > pick the "
              ".odt\n"
              "     4. confirm each source in the dialog; open the result -- "
              "citations are now live Zotero fields")

    if args.pdf:
        soffice = shutil.which("libreoffice") or shutil.which("soffice")
        if soffice:
            log("rendering PDF preview...")
            subprocess.run([soffice, "--headless", "--convert-to", "pdf",
                            "--outdir", os.path.dirname(out_path), out_path],
                           capture_output=True)
            print(f"OK  ->  {os.path.splitext(out_path)[0]}.pdf")
        else:
            log("LibreOffice not found; skipping --pdf")


if __name__ == "__main__":
    main()
