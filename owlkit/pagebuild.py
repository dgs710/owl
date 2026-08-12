"""Build the Word file one LaTeX page at a time.

The earlier approach tuned the whole document at once and inferred overflow
from the total page count.  That hides the thing you actually need to know:
*which* page is too full.  And it is misleading, because with hard page breaks
one overflowing page does not corrupt the content assignment -- it just
inserts an extra physical page, so every later page index is off by one and
the report looks far worse than the document is.

So this module treats each page as its own object, exactly as it is in the
PDF:

    page N  =  [float placed at top]  +  its prose  +  [float placed at bottom]

Every page is rendered **on its own**, in a document with the same styles and
geometry, and asked one question: does it fit on a single page?  Pages that do
not fit are re-measured at successively wider text blocks.  The widest
requirement across all pages becomes the single uniform setting for the whole
document, which is what was asked for.
"""

import copy
import os
import shutil
import subprocess
import tempfile

from docx import Document
from docx.oxml.ns import qn

from . import floats, pagefit


# ---------------------------------------------------------------------------
# cutting the document into pages
# ---------------------------------------------------------------------------

def page_elements(docx_path, pdf_words, counts, placements, density):
    """Return {page_no: [xml elements]} plus the prepared Document.

    The document is split exactly once; the same split is reused for every
    measurement, so page contents never drift between rounds.
    """
    doc = Document(docx_path)
    pagefit.strip_existing_page_breaks(doc)
    pagefit.apply_density(doc, density)

    blocks = floats.find_blocks(doc)
    if density.image_width_mm:
        # 11 in page - 2 x margin, minus room for a caption, in millimetres
        usable_mm = (11.0 - 2 * density.margin_in) * 25.4
        pagefit.set_float_image_width(blocks, density.image_width_mm,
                                      max_height_mm=usable_mm * 0.80)
    floats.detach(blocks)

    words, sites = pagefit.docx_word_stream(doc)
    mapped = pagefit.align_boundaries(pdf_words, counts, words)

    prose_pages = {n: counts[n - 1] > 0 for n in range(1, len(counts) + 1)}
    starts = pagefit.apply_boundaries(doc, sites, mapped, prose_pages)

    body = doc.element.body
    children = [el for el in body if el.tag != qn("w:sectPr")]

    # index every prose page start
    start_index = {}
    for page_no, el in starts.items():
        try:
            start_index[page_no] = children.index(el)
        except ValueError:
            pass

    n_pages = len(counts)
    pages = {n: [] for n in range(1, n_pages + 1)}

    # prose: everything from this page's start up to the next page's start
    ordered = sorted(start_index.items(), key=lambda kv: kv[1])
    bounds = []
    prev_page, prev_idx = 1, 0
    for page_no, idx in ordered:
        bounds.append((prev_page, prev_idx, idx))
        prev_page, prev_idx = page_no, idx
    bounds.append((prev_page, prev_idx, len(children)))
    for page_no, lo, hi in bounds:
        pages.setdefault(page_no, []).extend(children[lo:hi])

    # floats go to the page the PDF put them on, at the right end of it
    for kind, num, group in blocks:
        target = placements.get((kind, num))
        if target is None:
            continue
        page_no, where = target[0] + 1, target[1]
        slot = pages.setdefault(page_no, [])
        if where == "top":
            pages[page_no] = list(group) + slot
        else:
            pages[page_no] = slot + list(group)

    return doc, pages, blocks


# ---------------------------------------------------------------------------
# measuring one page on its own
# ---------------------------------------------------------------------------

def _blank_like(base_path, density, extra_bottom_in=0.0):
    """An empty document carrying the same styles and geometry.

    `extra_bottom_in` shrinks the frame for measurement only.  It is the
    safety margin: the measurement runs in LibreOffice but the document is
    opened in Word, and the two lay text out slightly differently.  A page
    that only just fits here can spill there -- and a spilled page drags a
    blank one after it, which is exactly how a 36-page document came back as
    38.  Requiring the page to fit in a frame this much shorter guarantees the
    headroom.
    """
    from docx.shared import Twips
    doc = Document(base_path)
    pagefit.apply_density(doc, density)
    if extra_bottom_in:
        for section in doc.sections:
            section.bottom_margin = Twips(
                int((density.margin_in + extra_bottom_in) * 1440))
    body = doc.element.body
    for el in list(body):
        if el.tag != qn("w:sectPr"):
            body.remove(el)
    return doc


def write_page_probes(base_path, pages, density, outdir, only=None,
                      extra_bottom_in=0.0):
    """Write one small .docx per page.  Returns {page_no: path}."""
    paths = {}
    for page_no, els in sorted(pages.items()):
        if only is not None and page_no not in only:
            continue
        doc = _blank_like(base_path, density, extra_bottom_in)
        body = doc.element.body
        sect = body.find(qn("w:sectPr"))
        for el in els:
            clone = copy.deepcopy(el)
            # a probe must not carry a page break of its own, or it always
            # measures as two pages
            pPr = clone.find(qn("w:pPr"))
            if pPr is not None:
                for pbb in pPr.findall(qn("w:pageBreakBefore")):
                    pPr.remove(pbb)
                for sp in pPr.findall(qn("w:sectPr")):
                    pPr.remove(sp)
            for br in clone.iter(qn("w:br")):
                if br.get(qn("w:type")) == "page":
                    br.getparent().remove(br)
            if sect is not None:
                sect.addprevious(clone)
            else:
                body.append(clone)
        path = os.path.join(outdir, f"probe_{page_no:03d}.docx")
        doc.save(path)
        paths[page_no] = path
    return paths


def batch_render(paths, outdir, chunk=12):
    """Convert many .docx to PDF in as few LibreOffice runs as possible."""
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise RuntimeError("LibreOffice is required to measure page fit")
    files = list(paths)
    for i in range(0, len(files), chunk):
        subprocess.run([soffice, "--headless", "--convert-to", "pdf",
                        "--outdir", outdir] + files[i:i + chunk],
                       capture_output=True, timeout=900)
    return {f: os.path.join(outdir, os.path.splitext(os.path.basename(f))[0] + ".pdf")
            for f in files}


def fill_ratio(pdf_path, margin_in):
    """How much of the text frame the content actually occupies, 0-1.

    "It fits" is not a safe test.  The measurement runs in LibreOffice but the
    document is opened in Word, and the two lay text out slightly differently:
    a page measured at 98% full here can spill in Word and take a blank page
    with it.  Requiring real headroom is what makes the result survive the
    change of renderer.
    """
    import pymupdf
    doc = pymupdf.open(pdf_path)
    if not len(doc):
        return 1.0
    page = doc[0]
    h = page.rect.height
    margin = margin_in * 72.0
    bottom = None
    for block in page.get_text("dict")["blocks"]:
        y1 = block["bbox"][3]
        bottom = y1 if bottom is None else max(bottom, y1)
    for d in page.get_drawings():
        y0, y1 = d["rect"][1], d["rect"][3]
        if y1 - y0 < h * 0.98:
            bottom = y1 if bottom is None else max(bottom, y1)
    doc.close()
    if bottom is None:
        return 0.0
    usable = h - 2 * margin
    return max(0.0, (bottom - margin) / usable) if usable > 0 else 1.0


def measure(base_path, pages, density, workdir, only=None, safety_in=0.30):
    """{page_no: rendered_page_count} with a frame `safety_in` inches shorter."""
    probe_dir = tempfile.mkdtemp(prefix="probe_", dir=workdir)
    paths = write_page_probes(base_path, pages, density, probe_dir, only,
                              extra_bottom_in=safety_in)
    rendered = batch_render(list(paths.values()), probe_dir)
    out = {}
    for page_no, docx in paths.items():
        pdf = rendered.get(docx)
        out[page_no] = pagefit.page_count(pdf) if pdf and os.path.exists(pdf) else 0
    return out


def overflowing(measurements):
    """Page numbers that do not fit inside the safety frame."""
    return sorted(n for n, pages in measurements.items() if pages != 1)


# ---------------------------------------------------------------------------
# assembling the final document
# ---------------------------------------------------------------------------

def assemble(base_path, pages, density, out_path, header_source=None):
    """Write the whole document out, page by page, in order.

    Each page's first element carries the page break, so the boundaries hold
    by construction rather than by hoping Word breaks in the same place.
    """
    doc = _blank_like(base_path, density)
    body = doc.element.body
    sect = body.find(qn("w:sectPr"))

    placed = 0
    # A <w:sectPr> inside a paragraph ends a section, and a section break of
    # type nextPage IS a page break.  The postprocessor puts them in so the
    # appendix can renumber its pages.  Adding our own pageBreakBefore on the
    # paragraph after one gives Word two breaks and an empty page in between,
    # so the previous page's section break is allowed to do the job instead.
    section_break_pending = False

    for page_no in sorted(pages):
        els = pages[page_no]
        first = True
        for el in els:
            clone = copy.deepcopy(el)
            pPr = clone.find(qn("w:pPr"))
            if pPr is not None:
                for pbb in pPr.findall(qn("w:pageBreakBefore")):
                    pPr.remove(pbb)
            for br in list(clone.iter(qn("w:br"))):
                if br.get(qn("w:type")) == "page":
                    br.getparent().remove(br)
            if first:
                if page_no > 1 and not section_break_pending:
                    pagefit._set_page_break_before(clone, True)
                section_break_pending = False
                first = False
            if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
                section_break_pending = True
            if sect is not None:
                sect.addprevious(clone)
            else:
                body.append(clone)
            placed += 1

    # keep the headers and footers the postprocessor built
    if header_source:
        src_doc = Document(header_source)
        for i, section in enumerate(doc.sections):
            if i < len(src_doc.sections):
                pass  # section properties already came from the same base file

    doc.save(out_path)
    return placed


# ---------------------------------------------------------------------------
# the whole page-matching run
# ---------------------------------------------------------------------------

# Loose to tight.  Tightening is monotonic -- a page that fits at one rung
# still fits at every later one -- so a page only ever has to be re-measured
# until it first passes.
DENSITY_LADDER = [
    (1.00, 140, 1.00),
    (0.95, 140, 1.00),
    (0.90, 140, 0.98),
    (0.85, 140, 0.97),
    (0.85, 135, 0.96),
    (0.80, 135, 0.96),
    (0.75, 130, 0.95),
    (0.70, 130, 0.95),
    (0.65, 125, 0.94),
    (0.60, 120, 0.93),
    (0.55, 115, 0.92),
    (0.50, 110, 0.90),
]


def match_pages(docx_path, pdf_path, tex_src, out_path,
                on_progress=None, safety_in=1.0, ladder=None):
    """Rebuild the .docx so its pages break where the compiled PDF's do.

    `on_progress(fraction, message)` is called throughout; the search is the
    slow part, so it reports per rung of the ladder.

    Returns a dict describing what it settled on and how well it matched.
    """
    from . import pagefit

    def say(frac, msg):
        if on_progress:
            on_progress(max(0.0, min(1.0, frac)), msg)

    ladder = ladder or DENSITY_LADDER

    say(0.02, "reading the compiled PDF")
    placements = floats.placements(pdf_path)
    captions = pagefit.caption_index(tex_src)
    pdf_words, counts = pagefit.page_word_counts(pdf_path, placements, captions)
    target_pages = len(counts)
    say(0.08, f"{target_pages} pages in the PDF, {len(placements)} figures and tables")

    workdir = tempfile.mkdtemp(prefix="owlmatch_")
    chosen = None
    pending = None                      # pages still to prove; None = all

    for i, (margin, mm, spacing) in enumerate(ladder):
        d = pagefit.Density(margin_in=margin, image_width_mm=mm,
                            line_spacing=spacing)
        frac = 0.10 + 0.65 * (i / max(1, len(ladder)))
        say(frac, f"fitting pages — margins {margin:.2f} in, figures {mm} mm")

        doc, pages, blocks = page_elements(docx_path, pdf_words, counts,
                                           placements, d)
        result = measure(docx_path, pages, d, workdir, only=pending,
                         safety_in=safety_in)
        bad = set(overflowing(result))
        if not bad:
            chosen = d
            say(frac + 0.03, f"every page fits with {safety_in:.1f} in to spare")
            break
        pending = bad                   # only these still need proving

    if chosen is None:
        chosen = pagefit.Density(margin_in=ladder[-1][0],
                                 image_width_mm=ladder[-1][1],
                                 line_spacing=ladder[-1][2])
        say(0.75, f"{len(pending)} page(s) remain tight even at the "
                  f"smallest margins")

    say(0.80, "assembling the document")
    doc, pages, blocks = page_elements(docx_path, pdf_words, counts,
                                       placements, chosen)
    assemble(docx_path, pages, chosen, out_path)

    say(0.88, "checking the result against the PDF")
    rendered = pagefit.render_pdf(out_path, workdir)
    got_pages = pagefit.page_count(rendered)
    checks = pagefit.verify(pdf_path, rendered, placements, captions)
    exact = sum(1 for c in checks if c.first_ok)

    say(1.0, f"{got_pages} pages, {exact}/{len(checks)} starting on the "
             f"same word as the PDF")
    return {
        "density": chosen,
        "pages": got_pages,
        "target_pages": target_pages,
        "exact_starts": exact,
        "checks": checks,
        "unresolved": sorted(pending) if chosen is None else [],
    }
